#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Script para gerar relatório de Projeto de Extensão III em formato Word
seguindo a norma NBR-15287:2025
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_page_break(doc):
    """Adiciona quebra de página"""
    doc.add_page_break()

def set_margins(doc, top=3, right=2, bottom=2, left=3):
    """Define as margens do documento (em cm)"""
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(top / 2.54)
        section.right_margin = Inches(right / 2.54)
        section.bottom_margin = Inches(bottom / 2.54)
        section.left_margin = Inches(left / 2.54)

def set_line_spacing(paragraph, spacing=1.5):
    """Define o espaçamento de linhas"""
    paragraph_format = paragraph.paragraph_format
    paragraph_format.line_spacing = spacing

def add_heading_section(doc, level, text):
    """Adiciona um heading como seção"""
    heading = doc.add_heading(text, level=level)
    heading.style = f'Heading {level}'
    set_line_spacing(heading, 1.5)
    return heading

def create_titulo_page(doc):
    """Cria página de título"""
    # Centralizando e adicionando espaços
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Centro Universitário
    p = doc.add_paragraph("Centro Universitário União das Américas Descomplica")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_line_spacing(p, 1.5)
    
    # Curso
    p = doc.add_paragraph("Curso de Tecnologia em Ciência de Dados")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_line_spacing(p, 1.5)
    
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Título do relatório
    p = doc.add_paragraph("RELATÓRIO DE PROJETO DE EXTENSÃO III")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(12)
    set_line_spacing(p, 1.5)
    
    # Subtítulo
    p = doc.add_paragraph("Análise da Situação - Ciência de Dados")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_line_spacing(p, 1.5)
    
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Informações finais
    p = doc.add_paragraph("Dezembro de 2025")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_line_spacing(p, 1.5)
    
    add_page_break(doc)

def create_sumario(doc):
    """Cria página de sumário"""
    p = doc.add_paragraph("SUMÁRIO")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    set_line_spacing(p, 1.5)
    
    doc.add_paragraph()
    
    items = [
        "1 INTRODUÇÃO E APRESENTAÇÃO",
        "2 ANÁLISE CRÍTICA DA SITUAÇÃO-PROBLEMA",
        "2.1 Identificação dos Problemas",
        "2.2 Aplicação de Técnicas de Análise",
        "3 IDENTIFICAÇÃO DOS FATORES-CHAVE E SOLUÇÃO",
        "3.1 Fatores Críticos Identificados",
        "3.2 Solução Desenvolvida",
        "3.3 Integração Entre Sistemas",
        "4 TECNOLOGIAS E ARQUITETURA",
        "4.1 Stack Tecnológico",
        "4.2 Estrutura do Projeto",
        "4.3 Modelo de Dados",
        "5 COMPETÊNCIAS DESENVOLVIDAS",
        "5.1 Competências Técnicas",
        "5.2 Soft Skills Desenvolvidas",
        "6 TEMAS ENVOLVIDOS NO PROJETO",
        "7 RESULTADOS E CONCLUSÃO",
        "7.1 Resultados Alcançados",
        "7.2 Benefícios para a Organização",
        "7.3 Conclusão",
        "REFERÊNCIAS",
    ]
    
    for item in items:
        p = doc.add_paragraph(item, style='List Bullet')
        set_line_spacing(p, 1.5)
    
    add_page_break(doc)

def add_introduction(doc):
    """Adiciona seção de introdução"""
    add_heading_section(doc, 1, "1 INTRODUÇÃO E APRESENTAÇÃO")
    
    p = doc.add_paragraph(
        "Iniciei meu projeto preenchendo a CARTA DE APRESENTAÇÃO e escolhi a organização "
        "XXXXXXXXXXXXXX para visitar e me apresentar. Após a autorização da organização, com o "
        "preenchimento do TERMO DE AUTORIZAÇÃO PARA REALIZAÇÃO DAS ATIVIDADES EXTENSIONISTAS, "
        "dei início ao projeto."
    )
    set_line_spacing(p, 1.5)
    
    # Sobre o projeto
    p = doc.add_paragraph()
    p_title = doc.add_paragraph("Sobre o Projeto")
    p_title.runs[0].bold = True
    set_line_spacing(p_title, 1.5)
    
    p = doc.add_paragraph(
        "Este projeto de extensão me proporcionou a oportunidade de me envolver diretamente com o "
        "uso de dados para resolver problemas reais enfrentados pela organização parceira. O foco "
        "foi na coleta, análise e interpretação de dados, visando a tomada de decisões mais informada "
        "e eficiente."
    )
    set_line_spacing(p, 1.5)
    
    # Objetivos
    p_title = doc.add_paragraph("Objetivos do Projeto")
    p_title.runs[0].bold = True
    set_line_spacing(p_title, 1.5)
    
    objetivos = [
        "Realizar uma análise detalhada dos problemas e necessidades da instituição",
        "Identificar questões que possam ser abordadas com técnicas de ciência de dados",
        "Aplicar modelagem preditiva, visualização e análise de dados",
        "Propor intervenções baseadas em insights provenientes da análise",
        "Conectar teoria e prática através de soluções com impacto positivo"
    ]
    
    for obj in objetivos:
        p = doc.add_paragraph(obj, style='List Bullet')
        set_line_spacing(p, 1.5)

def add_analise_critica(doc):
    """Adiciona seção de análise crítica"""
    add_heading_section(doc, 1, "2 ANÁLISE CRÍTICA DA SITUAÇÃO-PROBLEMA")
    
    # Identificação dos problemas
    add_heading_section(doc, 2, "2.1 Identificação dos Problemas")
    
    p = doc.add_paragraph(
        "Identifiquei e documentei os seguintes problemas na instituição, com foco na análise de dados:"
    )
    set_line_spacing(p, 1.5)
    
    # Problema 1
    p = doc.add_paragraph("Problema 1: Falta de Controle Financeiro")
    p.runs[0].bold = True
    set_line_spacing(p, 1.5)
    
    p = doc.add_paragraph(
        "A organização não possuía um sistema adequado para gerenciar suas finanças. Os registros de "
        "receitas e despesas eram feitos de forma manual ou em planilhas desorganizadas, dificultando "
        "a visualização do fluxo de caixa e a tomada de decisões financeiras."
    )
    set_line_spacing(p, 1.5)
    
    p = doc.add_paragraph("Impactos identificados:")
    p.runs[0].bold = True
    set_line_spacing(p, 1.5)
    
    impactos = [
        "Ausência de categorização de gastos",
        "Dificuldade em identificar padrões de despesas",
        "Impossibilidade de gerar relatórios analíticos",
        "Falta de visibilidade sobre a saúde financeira"
    ]
    
    for imp in impactos:
        p = doc.add_paragraph(imp, style='List Bullet')
        set_line_spacing(p, 1.5)
    
    # Problema 2
    p = doc.add_paragraph("Problema 2: Gestão de Estoque Ineficiente")
    p.runs[0].bold = True
    set_line_spacing(p, 1.5)
    
    p = doc.add_paragraph(
        "O controle de estoque era precário, resultando em perdas por falta de produtos ou excesso "
        "de itens parados. Não havia registro adequado de compras e vendas, nem integração com o "
        "controle financeiro."
    )
    set_line_spacing(p, 1.5)
    
    p = doc.add_paragraph("Impactos identificados:")
    p.runs[0].bold = True
    set_line_spacing(p, 1.5)
    
    impactos = [
        "Desconhecimento do saldo real de produtos",
        "Falta de histórico de movimentações",
        "Ausência de alertas para estoque baixo",
        "Dificuldade em calcular lucro por produto"
    ]
    
    for imp in impactos:
        p = doc.add_paragraph(imp, style='List Bullet')
        set_line_spacing(p, 1.5)
    
    # Aplicação de técnicas de análise
    add_heading_section(doc, 2, "2.2 Aplicação de Técnicas de Análise")
    
    p = doc.add_paragraph(
        "Apliquei técnicas de análise exploratória de dados para compreender a fundo as questões identificadas:"
    )
    set_line_spacing(p, 1.5)
    
    # Timeline
    timeline_items = [
        ("Coleta de Dados", 
         "Coletei informações sobre as transações financeiras, movimentações de estoque, produtos "
         "comercializados e fluxo de caixa histórico da organização."),
        ("Análise Exploratória",
         "Realizei análise exploratória para identificar padrões de gastos, sazonalidade nas vendas "
         "e produtos com maior rotatividade."),
        ("Identificação de Padrões",
         "Identifiquei padrões de comportamento nos dados que indicavam oportunidades de melhoria nos processos.")
    ]
    
    for title, desc in timeline_items:
        p = doc.add_paragraph(f"{title}", style='List Number')
        set_line_spacing(p, 1.5)
        p_desc = doc.add_paragraph(desc)
        set_line_spacing(p_desc, 1.5)

def add_fatores_solucao(doc):
    """Adiciona seção de fatores-chave e solução"""
    add_heading_section(doc, 1, "3 IDENTIFICAÇÃO DOS FATORES-CHAVE E SOLUÇÃO")
    
    # Fatores críticos
    add_heading_section(doc, 2, "3.1 Fatores Críticos Identificados")
    
    p = doc.add_paragraph(
        "Identifiquei as variáveis e fatores críticos que influenciam os problemas, utilizando "
        "metodologias de análise:"
    )
    set_line_spacing(p, 1.5)
    
    # Tabela de fatores
    table = doc.add_table(rows=6, cols=3)
    table.style = 'Light Grid Accent 1'
    
    # Cabeçalho
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Fator'
    hdr_cells[1].text = 'Impacto'
    hdr_cells[2].text = 'Solução Proposta'
    
    fatores = [
        ("Registro manual de transações", "Alto - Erros e perda de dados", "Sistema digital automatizado"),
        ("Falta de categorização", "Médio - Análise prejudicada", "Categorias personalizáveis"),
        ("Ausência de relatórios", "Alto - Decisões sem dados", "Dashboard analítico"),
        ("Estoque descontrolado", "Alto - Perdas financeiras", "Sistema de gestão integrado"),
        ("Sem integração financeira", "Médio - Visão fragmentada", "Integração automática"),
    ]
    
    for i, (fator, impacto, solucao) in enumerate(fatores, 1):
        row_cells = table.rows[i].cells
        row_cells[0].text = fator
        row_cells[1].text = impacto
        row_cells[2].text = solucao
    
    doc.add_paragraph()
    
    # Solução desenvolvida
    add_heading_section(doc, 2, "3.2 Solução Desenvolvida")
    
    p = doc.add_paragraph(
        "Desenvolvi um sistema integrado composto por dois módulos principais que se comunicam entre si:"
    )
    set_line_spacing(p, 1.5)
    
    # Sistema Financeiro
    p = doc.add_paragraph("Sistema Financeiro (Porta 5000)")
    p.runs[0].bold = True
    set_line_spacing(p, 1.5)
    
    p = doc.add_paragraph("Sistema completo para gestão financeira com:")
    set_line_spacing(p, 1.5)
    
    features = [
        "Dashboard com visão geral de receitas e despesas",
        "Cadastro de lançamentos (contas a pagar/receber)",
        "Categorização personalizável",
        "Filtros por período",
        "Relatórios analíticos",
        "Formas de pagamento configuráveis"
    ]
    
    for feat in features:
        p = doc.add_paragraph(feat, style='List Bullet')
        set_line_spacing(p, 1.5)
    
    # Sistema de Estoque
    p = doc.add_paragraph("Sistema de Estoque (Porta 5001)")
    p.runs[0].bold = True
    set_line_spacing(p, 1.5)
    
    p = doc.add_paragraph("Sistema completo para controle de estoque com:")
    set_line_spacing(p, 1.5)
    
    features = [
        "Cadastro de produtos com código de barras",
        "Controle de compras e vendas",
        "Ajuste de estoque por inventário",
        "Cálculo automático de lucro",
        "Alertas de estoque baixo",
        "Integração com sistema financeiro"
    ]
    
    for feat in features:
        p = doc.add_paragraph(feat, style='List Bullet')
        set_line_spacing(p, 1.5)
    
    # Integração
    add_heading_section(doc, 2, "3.3 Integração Entre Sistemas")
    
    p = doc.add_paragraph(
        "A principal inovação é a integração automática entre os sistemas. Quando uma compra é "
        "registrada no sistema de estoque, automaticamente é gerado um lançamento de despesa no "
        "sistema financeiro. Da mesma forma, quando uma venda é realizada, é gerado um lançamento "
        "de receita (contas a receber). Isso garante consistência dos dados e visão unificada da "
        "saúde financeira do negócio."
    )
    set_line_spacing(p, 1.5)

def add_tecnologias(doc):
    """Adiciona seção de tecnologias"""
    add_heading_section(doc, 1, "4 TECNOLOGIAS E ARQUITETURA")
    
    # Stack tecnológico
    add_heading_section(doc, 2, "4.1 Stack Tecnológico")
    
    techs = [
        "Python 3.x",
        "Flask",
        "HTML5",
        "CSS3",
        "JavaScript",
        "Bootstrap 5",
        "JSON",
        "CSV"
    ]
    
    for tech in techs:
        p = doc.add_paragraph(tech, style='List Bullet')
        set_line_spacing(p, 1.5)
    
    # Estrutura
    add_heading_section(doc, 2, "4.2 Estrutura do Projeto")
    
    estrutura = """pex3/
├── financeiro.py                # Sistema Financeiro (Flask)
├── estoque.py                   # Sistema de Estoque (Flask)
├── iniciar_sistemas.py          # Script para iniciar ambos
├── database.json                # Banco de dados financeiro
├── estoque_db.json              # Banco de dados de movimentações
├── produtos.csv                 # Cadastro de produtos
└── templates/
    ├── base.html                # Template base financeiro
    ├── index.html               # Dashboard financeiro
    ├── lancamentos.html         # Lista de lançamentos
    ├── form_lancamento.html     # Formulário de lançamento
    ├── categorias.html          # Gestão de categorias
    ├── analytics.html           # Relatórios analíticos
    └── estoque/
        ├── base_estoque.html    # Template base estoque
        ├── index.html           # Dashboard estoque
        ├── produtos.html        # Lista de produtos
        ├── form_produto.html    # Cadastro de produto
        ├── compras.html         # Lista de compras
        ├── form_compra.html     # Nova compra
        ├── vendas.html          # Lista de vendas
        ├── form_venda.html      # Nova venda
        └── relatorios.html      # Relatórios"""
    
    p = doc.add_paragraph(estrutura)
    p.style = 'No Spacing'
    p.runs[0].font.name = 'Courier New'
    p.runs[0].font.size = Pt(10)
    set_line_spacing(p, 1.0)
    
    # Modelo de dados
    add_heading_section(doc, 2, "4.3 Modelo de Dados")
    
    p = doc.add_paragraph("database.json (Financeiro)")
    p.runs[0].bold = True
    set_line_spacing(p, 1.5)
    
    json_example = """{
  "transactions": [
    {
      "id": 1,
      "tipo": "receber",
      "data_gasto": "2025-12-30",
      "valor": 200.0,
      "categoria": "Venda",
      "forma_pagamento": "PIX",
      "descricao": "..."
    }
  ],
  "categories": [...],
  "payment_methods": [...]
}"""
    
    p = doc.add_paragraph(json_example)
    p.style = 'No Spacing'
    p.runs[0].font.name = 'Courier New'
    p.runs[0].font.size = Pt(10)
    set_line_spacing(p, 1.0)
    
    p = doc.add_paragraph("produtos.csv (Estoque)")
    p.runs[0].bold = True
    set_line_spacing(p, 1.5)
    
    csv_example = """codigo_barras;nome;saldo;preco_venda;preco_compra
7891234567890;Arroz;50;12.90;8.50
7891234567891;Feijão;30;9.90;6.20"""
    
    p = doc.add_paragraph(csv_example)
    p.style = 'No Spacing'
    p.runs[0].font.name = 'Courier New'
    p.runs[0].font.size = Pt(10)
    set_line_spacing(p, 1.0)

def add_competencias(doc):
    """Adiciona seção de competências"""
    add_heading_section(doc, 1, "5 COMPETÊNCIAS DESENVOLVIDAS")
    
    # Competências técnicas
    add_heading_section(doc, 2, "5.1 Competências Técnicas")
    
    tecnicas = [
        ("Análise de Dados e Estatística",
         "Desenvolvi a capacidade de compreender, interpretar e manipular dados em diferentes "
         "contextos, aplicando técnicas estatísticas para extrair insights relevantes."),
        ("Modelagem e Algoritmos",
         "Adquiri competência em criar modelos de dados e estruturas algorítmicas para resolver "
         "problemas específicos de negócio."),
        ("Visualização de Dados",
         "Desenvolvi habilidade para comunicar insights de forma visual e acessível, por meio de "
         "dashboards e relatórios interativos."),
        ("Ferramentas de Ciência de Dados",
         "Obtive conhecimento prático em Python, Flask, JSON, CSV e bibliotecas para desenvolvimento "
         "de aplicações orientadas a dados.")
    ]
    
    for title, desc in tecnicas:
        p = doc.add_paragraph(f"{title}")
        p.runs[0].bold = True
        set_line_spacing(p, 1.5)
        p = doc.add_paragraph(desc)
        set_line_spacing(p, 1.5)
    
    # Soft skills
    add_heading_section(doc, 2, "5.2 Soft Skills Desenvolvidas")
    
    skills = [
        ("Pensamento Analítico",
         "Desenvolvi habilidade para compreender problemas complexos e propor soluções baseadas em dados e evidências."),
        ("Trabalho em Equipe",
         "Exercitei a capacidade de colaborar com membros da organização para entender suas necessidades de dados."),
        ("Iniciativa e Autonomia",
         "Demonstrei proatividade e independência na coleta, análise de dados e desenvolvimento da solução."),
        ("Comunicação Eficaz",
         "Aprimorei a competência em transmitir insights de dados de maneira clara e compreensível para não-técnicos."),
        ("Flexibilidade",
         "Desenvolvi habilidade para me ajustar a diferentes ferramentas, contextos e requisitos durante o projeto."),
        ("Responsabilidade Ética",
         "Mantive compromisso com o uso responsável de dados, respeitando privacidade e boas práticas.")
    ]
    
    for title, desc in skills:
        p = doc.add_paragraph(f"{title}")
        p.runs[0].bold = True
        set_line_spacing(p, 1.5)
        p = doc.add_paragraph(desc)
        set_line_spacing(p, 1.5)

def add_temas(doc):
    """Adiciona seção de temas envolvidos"""
    add_heading_section(doc, 1, "6 TEMAS ENVOLVIDOS NO PROJETO")
    
    temas = [
        ("Exploração e Limpeza de Dados",
         "Apliquei técnicas para organizar, limpar e preparar dados para análise, garantindo a "
         "qualidade das informações processadas pelo sistema."),
        ("Visualização de Dados",
         "Utilizei ferramentas e técnicas para criar visualizações eficazes que comunicam resultados "
         "de forma clara nos dashboards."),
        ("Estatística Aplicada",
         "Empreguei conceitos estatísticos para interpretar dados financeiros e de estoque, apoiando "
         "a tomada de decisão."),
        ("Análise de Dados para Negócios",
         "Desenvolvi soluções de análise de dados para otimização de processos e suporte a decisões "
         "estratégicas da organização.")
    ]
    
    for title, desc in temas:
        p = doc.add_paragraph(f"{title}")
        p.runs[0].bold = True
        set_line_spacing(p, 1.5)
        p = doc.add_paragraph(desc)
        set_line_spacing(p, 1.5)

def add_resultados(doc):
    """Adiciona seção de resultados"""
    add_heading_section(doc, 1, "7 RESULTADOS E CONCLUSÃO")
    
    # Resultados alcançados
    add_heading_section(doc, 2, "7.1 Resultados Alcançados")
    
    p = doc.add_paragraph("Entregas do Projeto:")
    p.runs[0].bold = True
    set_line_spacing(p, 1.5)
    
    entregas = [
        "Sistema Financeiro Completo: Aplicação web funcional para gestão de finanças",
        "Sistema de Estoque Integrado: Controle completo de produtos, compras e vendas",
        "Integração Automática: Comunicação entre sistemas para consistência de dados",
        "Dashboards Analíticos: Visualizações para tomada de decisão",
        "Documentação: Código comentado e relatório detalhado"
    ]
    
    for entrega in entregas:
        p = doc.add_paragraph(entrega, style='List Bullet')
        set_line_spacing(p, 1.5)
    
    # Benefícios
    add_heading_section(doc, 2, "7.2 Benefícios para a Organização")
    
    table = doc.add_table(rows=6, cols=2)
    table.style = 'Light Grid Accent 1'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Antes'
    hdr_cells[1].text = 'Depois'
    
    beneficios = [
        ("Controle manual em planilhas", "Sistema digital automatizado"),
        ("Dados dispersos e inconsistentes", "Base de dados centralizada e integrada"),
        ("Sem visibilidade financeira", "Dashboard com visão em tempo real"),
        ("Estoque descontrolado", "Controle preciso com alertas"),
        ("Decisões sem dados", "Decisões baseadas em dados"),
    ]
    
    for i, (antes, depois) in enumerate(beneficios, 1):
        row_cells = table.rows[i].cells
        row_cells[0].text = antes
        row_cells[1].text = depois
    
    doc.add_paragraph()
    
    # Conclusão
    add_heading_section(doc, 2, "7.3 Conclusão")
    
    p = doc.add_paragraph(
        "Este projeto de extensão me permitiu aplicar na prática os conhecimentos adquiridos no "
        "curso de Ciência de Dados, desenvolvendo uma solução real que traz impacto positivo para "
        "a organização parceira. A experiência de identificar problemas, analisar dados e propor "
        "soluções baseadas em evidências foi fundamental para meu desenvolvimento profissional."
    )
    set_line_spacing(p, 1.5)
    
    p = doc.add_paragraph(
        "O sistema desenvolvido resolve problemas concretos de gestão financeira e controle de "
        "estoque, automatizando processos que antes eram manuais e propensos a erros. A integração "
        "entre os módulos garante consistência dos dados e oferece uma visão unificada da saúde do negócio."
    )
    set_line_spacing(p, 1.5)

def add_referencias(doc):
    """Adiciona seção de referências"""
    add_heading_section(doc, 1, "REFERÊNCIAS")
    
    referencias = [
        "ASSUNÇÃO, R. M., & OLIVEIRA, J. P. (2016). Inclusão digital e alfabetização tecnológica: um estudo de caso. Salvador: EDUFBA.",
        
        "BATISTA, E. S. (2012). Tecnologias assistivas e inclusão digital. São Paulo: Cultura Acadêmica.",
        
        "KEEGAN, V. (2015). Desenvolvimento de jogos digitais. São Paulo: Novatec.",
        
        "MENDES, C. L. (2018). Segurança da informação: uma visão gerencial. São Paulo: Saraiva.",
        
        "MONTEIRO, M. (2014). Design para a Internet: projetando a experiência perfeita. Rio de Janeiro: Alta Books.",
        
        "NORTON, P. (2002). Introdução à informática. São Paulo: Makron Books.",
        
        "NUNES, C. S. (2017). Robótica educacional: princípios e práticas. Porto Alegre: Bookman.",
        
        "PEREIRA, J. R. M., & MENDES, L. F. (2015). Hackathons: inovando com maratonas de programação. São Paulo: Editora Blucher.",
        
        "PRESSMAN, R. S. (2019). Engenharia de software: uma abordagem profissional. 8. ed. Porto Alegre: AMGH.",
        
        "RIBEIRO, M. A., & ALVES, T. M. (2019). Sustentabilidade e tecnologia: estratégias e práticas. Rio de Janeiro: Elsevier.",
        
        "SOMMERVILLE, I. (2011). Engenharia de Software. 9. ed. São Paulo: Pearson.",
        
        "TANENBAUM, A. S., & WETHERALL, D. J. (2011). Redes de computadores. 5. ed. São Paulo: Pearson.",
    ]
    
    for ref in referencias:
        p = doc.add_paragraph(ref, style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)
        set_line_spacing(p, 1.0)

def main():
    """Função principal"""
    # Criar documento
    doc = Document()
    
    # Configurar margens (NBR-15287: esquerda 3cm, superior 3cm, direita 2cm, inferior 2cm)
    set_margins(doc, top=3, right=2, bottom=2, left=3)
    
    # Configurar fonte padrão
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Criar página de título
    create_titulo_page(doc)
    
    # Criar sumário
    create_sumario(doc)
    
    # Adicionar seções
    add_introduction(doc)
    add_analise_critica(doc)
    add_fatores_solucao(doc)
    add_tecnologias(doc)
    add_competencias(doc)
    add_temas(doc)
    add_resultados(doc)
    add_referencias(doc)
    
    # Adicionar rodapé com copyright
    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    p.text = "Copyright © Delean Mafra - Todos os direitos reservados | Licença: CC BY-NC 4.0"
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(10)
    
    # Salvar documento
    output_path = r"d:\Python\complementos\pex3\Relatorio_Extensao_NBR15287_2025.docx"
    doc.save(output_path)
    print(f"Documento criado com sucesso: {output_path}")

if __name__ == "__main__":
    main()
